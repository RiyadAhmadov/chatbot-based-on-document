import streamlit as st
import os
import logging
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_ollama import ChatOllama
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
import ollama

# =======================
# ⚙️ Configurations
# =======================
logging.basicConfig(level=logging.INFO)

DOC_PATH = "./test_doc.docx"
MODEL_NAME = "llama3.2"
PERSIST_DIRECTORY = "./chroma_db"

# =======================
# 🧩 Load DOCX
# =======================
def ingest_docx(doc_path):
    """Load DOCX as LangChain Documents."""
    if os.path.exists(doc_path):
        doc = DocxDocument(doc_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])

        # Extract tables
        tables_text = ""
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                tables_text += " | ".join(row_data) + "\n"
            tables_text += "\n"

        combined_text = full_text + "\n" + tables_text
        logging.info("✅ DOCX loaded successfully.")
        return [Document(page_content=combined_text)]
    else:
        logging.error(f"❌ DOCX file not found at path: {doc_path}")
        st.error("DOCX file not found.")
        return None

# =======================
# ✂️ Split Text
# =======================
def split_documents(documents):
    """Split documents into smaller chunks."""
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=300)
    chunks = text_splitter.split_documents(documents)
    logging.info(f"📚 Documents split into {len(chunks)} chunks.")
    return chunks

# =======================
# 💾 Vector DB
# =======================
@st.cache_resource
def load_vector_db():
    """Load or create the vector database (cached)."""
    embedding = HuggingFaceEmbeddings(model_name="paraphrase-multilingual-MiniLM-L12-v2")

    if os.path.exists(PERSIST_DIRECTORY) and os.listdir(PERSIST_DIRECTORY):
        vector_db = Chroma(
            embedding_function=embedding,
            persist_directory=PERSIST_DIRECTORY,
        )
        logging.info("🟢 Loaded existing Chroma DB.")
    else:
        data = ingest_docx(DOC_PATH)
        if data is None:
            return None
        chunks = split_documents(data)
        vector_db = Chroma.from_documents(
            documents=chunks,
            embedding=embedding,
            persist_directory=PERSIST_DIRECTORY,
        )
        vector_db.persist()
        logging.info("🆕 Vector database created and persisted.")
    return vector_db

# =======================
# 🔗 Create Chain
# =======================
def create_chain(vector_db, llm):
    """RAG zəncirini Azərbaycan dili üçün yaradın."""
    template = """
        Aşağıdakı kontekstdən istifadə edərək **sadə və aydın Azərbaycan dilində** suala cavab ver:
        {context}

        Sual: {question}

        Cavabı qısa, düzgün qrammatik və aydın yaz.
    """
    prompt = ChatPromptTemplate.from_template(template)
    retriever = vector_db.as_retriever(search_kwargs={"k": 3})

    chain = (
        {"context": retriever, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )
    return chain


# =======================
# 🎯 Streamlit App
# =======================
def main():
    st.title("⚡ Goldenpay DOCX Assistant (Fast Version)")
    user_input = st.text_input("Enter your question:", "")

    if user_input:
        with st.spinner("Generating response..."):
            try:
                # Initialize LLM (Mistral model from Ollama)
                llm = ChatOllama(model=MODEL_NAME)

                # Load cached vector DB
                vector_db = load_vector_db()
                if vector_db is None:
                    st.error("Failed to load or create the vector database.")
                    return

                # Build chain
                chain = create_chain(vector_db, llm)

                st.markdown("### 💬 Assistant:")
                response_stream = chain.stream(input=user_input)

                # Stream response token by token
                full_response = ""
                response_placeholder = st.empty()
                for chunk in response_stream:
                    full_response += chunk
                    response_placeholder.markdown(full_response)

                logging.info("✅ Response generated successfully.")
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
                logging.error(f"Error details: {e}")
    else:
        st.info("Please enter a question to get started.")

# =======================
if __name__ == "__main__":
    main()
